from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import newspaper
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
from textblob import TextBlob
from serpapi import GoogleSearch
import regex as re
from urllib.request import urlretrieve
from urllib.parse import urlencode
from youtube_transcript_api import YouTubeTranscriptApi

import lang

import config
from wordcloud import WordCloud
import matplotlib.pyplot as plt


data_wc = ""

def wc(keyword,text):
    wordcloud = WordCloud(background_color="white").generate(text)
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis("off")
    plt.savefig(f'./wc/{keyword}.jpeg')


def getSnapshot(link, id, type, keyword):
    id = str(keyword)+"_"+str(type)+"_"+str(id)
    filename = "./ss/" + id + ".jpeg"
    print("taking snap")
    params = urlencode(dict(access_key=config.APIFLASH,
                        url=link
                        
                        ))
    try:
        urlretrieve("https://api.apiflash.com/v1/urltoimage?" + params, filename )
    except KeyboardInterrupt:
        exit(0)
    return filename

def is_biography_page(url):
    biography_keywords = ["wiki", "biography", "profile", "about", "tag", "topic", "videos"]
    
    for keyword in biography_keywords:
        if keyword in url.split('/')[3].lower():
            return True
    
    return False

def is_socials(url):
    
    socials = ['instagram', 'facebook', 'youtube', 'vimeo', 'inhabitat.com', 'warwickonline', 'amazon', 'flipkart', 'webstories']
    for key in socials:
        if key in url.lower():
            return True
        
    return False

def extract_news_content(url):
    global data_wc
    try:
        article = newspaper.Article(url)
        article.download()
        article.parse()
    except newspaper.ArticleException:
        print("cannot fetch news for the url: ", url)
        return "False"
    
    title = article.title.encode('utf-8')
    body = article.text.encode('utf-8')
    # date = article.publish_date
    
    print("News Fetched!")
    data_wc += body.decode('utf-8')
    return {'title' :title, 'body':body,
            # 'date':date
            }

def get_sentiment_score(text):
    blob = TextBlob(text)
    sentiment_score = blob.sentiment.polarity
    return sentiment_score

def get_video_info(video_id):
    youtube = build('youtube', 'v3', developerKey=config.youtubeAPI)
    
    try:
        response = youtube.videos().list(
            part='snippet',
            id=video_id
        ).execute()

        items = response.get('items', [])
        if items:
            video_info = items[0]
            print(video_info)
            title = video_info['snippet']['title']
            # thumbnail_url = video_info['snippet']['thumbnails']['default']['url']
            return title

    except HttpError as e:
        print(f'Error retrieving video info for video ID {video_id}: {e}')

    return None, None

def getYoutubeLinks(keyword, max):
    data = []
    params = {
            "api_key": config.serpAPI,
            "engine": "youtube",
            "search_query": keyword,
            "hl": "en",
            "gl": "in"
        }
    search = GoogleSearch(params)
    results = search.get_dict()
    
    res = results.get('video_results')
    for i in res:
        data.append({
            'link': i.get('link'),
            'title': i.get('title'),
            'description': i.get('description'),
            'published': i.get('published_date'),
            'thumbnail': i.get('thumbnail').get("static")
        })
        
    return data[:max]

def generate_docx(category, id, title, summary, sentiment, link, filename):

    id = str(category)+"_"+str(id)
    document = Document()
    
    document.add_paragraph(str(title))
    document.add_picture(filename, width=Inches(7), height=Inches(5))
    document.add_paragraph(link)
    document.add_paragraph(summary)
    document.add_paragraph(str(sentiment))
    document.save(f"./reports/{id}.docx")
    
def generateReport(data, keyword):
    document = Document()
    global data_wc
    for i in data:
        # if float(i['sentiment']) <= 0.0:
        print(i['summary'], i['sentiment'])
        i['id'] = str(i['type'])+"_"+str(i['id'])
        document.add_paragraph(str(i['title']))
        document.add_picture(i['filename'], width=Inches(7), height=Inches(5))
        document.add_paragraph(i['link'])
        document.add_paragraph(i['summary'])
        document.add_paragraph(str(i['sentiment']))
        document.add_page_break()
        # else: continue
    wc(keyword, data_wc)
    document.add_picture(f'./wc/{keyword}.jpeg', width=Inches(7), height=Inches(5))
    document.save(f"./reports/{keyword}.docx")
    
def generate_pdf(id, title, summary, sentiment, link, filename):
    # print("HEHEHEHE", summary)
    print("Generating PDF for News: ", id)
    pdf = FPDF()
    pdf.add_page()

    pdf.set_font('Arial', 'B', 16)

    pdf.cell(200, 10, "Title: " + title, align='L')

    pdf.set_font('Arial', '', 12)
    pdf.image(filename,  x=10, y=80, w=220, h=150 )
    pdf.add_page()

    pdf.set_font('Arial', 'B', 12)
    pdf.multi_cell(150, 10,  "Summary: ")
    pdf.set_font('Arial', '', 12)
    pdf.multi_cell(150, 10,  summary)
    
    pdf.set_font('Arial', 'B', 12)
    pdf.multi_cell(180, 5, "URL:")
    pdf.set_font('Arial', '', 12)
    pdf.multi_cell(180, 5, link)
    pdf.set_font('Arial', 'B', 12)
    pdf.multi_cell(180, 5, "Polarity Score:")
    pdf.set_font('Arial', '', 12)
    pdf.multi_cell(180, 5, sentiment)
    pdf.set_margins(5, 5, 5)

    # pdfname= title+'.pdf'
    pdf.output(f"./reports/{id}.pdf")
    print("PDF ", id, " Generated!")

def extractVideoId(link):
    regex = r"(?<=v=|v\/|vi=|vi\/|youtu.be\/|embed\/|\/v\/|\/e\/|watch\?v=|\?v=|\/embed\/|\/e\/|youtu.be\/|\/v\/|watch\?v=|embed\/)[^#\\?\\&]*"
    match = re.search(regex, link)
    if match:
        return match.group(0)
    else:
        return None

def getVideoTranscript(video_id):
    global data_wc
    transcript = ""
    try:
        data = YouTubeTranscriptApi.get_transcript(video_id, languages=['en-GB', 'en', 'hindi', 'kannada', 'telugu', 'tamil', 'malayalam', 'bengali'])
        print(f"Transcript for video ID {video_id}:")
        for segment in data:
            transcript += "\n" + segment['text']
    except Exception as e:
        print(f"Error occurred for video ID {video_id}: {str(e)}")

    data_wc += transcript
    return transcript

def getVideoTitle(video_id):
    title = get_video_info(video_id)
    return title

def getVideoDescription(video_id):

    global data_wc
    youtube = build('youtube', 'v3', developerKey=config.youtubeAPI)
    
    try:
        response = youtube.videos().list(
            part='snippet',
            id=video_id
        ).execute()

        items = response.get('items', [])
        if items:
            video_info = items[0]
            description = video_info['snippet']['description']
            data_wc += description
            return description

    except HttpError as e:
        print(f'Error retrieving video info for video ID {video_id}: {e}')

    return None, None

from langchain.indexes import VectorstoreIndexCreator
from langchain.text_splitter import CharacterTextSplitter

def cb(body):
    text_splitter = CharacterTextSplitter()

    texts = text_splitter.split_text()

    docs = [Document(page_content=t) for t in texts[:3]]

    index = VectorstoreIndexCreator().from_loaders([docs])

    query = "What is the document about?"
    print(index.query(query))
    
