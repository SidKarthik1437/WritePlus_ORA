import os
from docx import Document
import docx
# from docx.enum.table import W
from docx.shared import Pt
from docx.shared import Inches

folder_path = "./reports/"  # Replace with the actual folder path
data = []
# Iterate over the files in the folder
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        
        # Load the DOCX document
        doc = Document(file_path)
        
        # Extract the URL
        url = doc.paragraphs[2].text  # Assuming the URL is always the third paragraph (index 2)
        
        # Extract the sentiment score
        sentiment_score = doc.paragraphs[4].text # Assuming the sentiment score is always the last paragraph
        
        # Process the extracted data
        print("File:", filename)
        print("URL:", url)
        print("Sentiment Score:", sentiment_score)
        print("---")
        data.append({
            'url': url,
            'sentiment_score': sentiment_score
        })

new_doc = Document()
table = new_doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.autofit = True

# Add table headers
headers = ["SL No", "URL", "Sentiment Score"]
for idx, header in enumerate(headers):
    table.cell(0, idx).text = header

# Populate table with extracted data
for idx, data in enumerate(data):
    row = table.add_row().cells
    row[0].text = str(idx + 1)
    row[1].text = data['url']
    row[2].text = str(data['sentiment_score'])

# Save the new document
new_doc.save("./table.docx")